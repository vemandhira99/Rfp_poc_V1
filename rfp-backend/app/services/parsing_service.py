import os, json
from sqlalchemy.orm import Session
from app.models.rfp import RFPDocument, AuditLog
from app.models.sections import RFPSection

# Keywords to detect section types
SECTION_KEYWORDS = {
    "scope":       ["scope", "objective", "purpose", "overview"],
    "deadline":    ["deadline", "timeline", "schedule", "due date", "submission"],
    "budget":      ["budget", "cost", "price", "financial", "commercial"],
    "legal":       ["legal", "compliance", "terms", "conditions", "liability"],
    "sla":         ["sla", "service level", "performance", "availability"],
    "technical":   ["technical", "architecture", "infrastructure", "technology"],
    "eligibility": ["eligibility", "qualification", "criteria", "requirement"],
    "contact":     ["contact", "address", "email", "phone", "submit"],
}

def detect_tags(text: str) -> list:
    text_lower = text.lower()
    tags = []
    for tag, keywords in SECTION_KEYWORDS.items():
        if any(kw in text_lower for kw in keywords):
            tags.append(tag)
    return tags

def parse_pdf(file_path: str) -> tuple:
    """Extract sections from PDF file"""
    sections = []
    total_pages = 0

    try:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)

            current_section = None
            current_text = []
            current_page = 1

            for page_num in range(total_pages):
                page = reader.pages[page_num]
                text = page.extract_text() or ""
                lines = text.split("\n")

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    # Detect if this line is a heading
                    is_heading = (
                        len(line) < 80 and
                        (line.isupper() or
                         line.endswith(":") or
                         any(line.lower().startswith(kw) for kws in SECTION_KEYWORDS.values() for kw in kws))
                    )

                    if is_heading and current_text:
                        # Save previous section
                        section_text = " ".join(current_text).strip()
                        if section_text and current_section:
                            sections.append({
                                "section_name": current_section,
                                "section_text": section_text[:5000],
                                "page_number": current_page,
                                "confidence": 85.0,
                                "tags": detect_tags(section_text)
                            })
                        current_section = line
                        current_text = []
                        current_page = page_num + 1
                    else:
                        if not current_section:
                            current_section = "Introduction"
                        current_text.append(line)

            # Save last section
            if current_text and current_section:
                section_text = " ".join(current_text).strip()
                if section_text:
                    sections.append({
                        "section_name": current_section or "Introduction",
                        "section_text": section_text[:5000],
                        "page_number": current_page,
                        "confidence": 85.0,
                        "tags": detect_tags(section_text)
                    })

        return sections, total_pages

    except Exception as e:
        print(f"PDF Parse Error: {e}")
        return [], 0


def parse_docx(file_path: str) -> tuple:
    """Extract sections from DOCX file"""
    sections = []
    try:
        from docx import Document
        doc = Document(file_path)
        page_count = max(1, len(doc.paragraphs) // 20) # Rough estimate
        current_section = "Introduction"
        current_text = []
        page_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            is_heading = para.style.name.startswith("Heading") or (
                len(text) < 80 and text.isupper()
            )

            if is_heading and current_text:
                section_text = " ".join(current_text).strip()
                if section_text:
                    sections.append({
                        "section_name": current_section,
                        "section_text": section_text[:5000],
                        "page_number": page_num,
                        "confidence": 90.0,
                        "tags": detect_tags(section_text)
                    })
                current_section = text
                current_text = []
            else:
                current_text.append(text)

        if current_text:
            section_text = " ".join(current_text).strip()
            sections.append({
                "section_name": current_section,
                "section_text": section_text[:5000],
                "page_number": page_num,
                "confidence": 90.0,
                "tags": detect_tags(section_text)
            })

        return sections, page_count
    except Exception as e:
        print(f"DOCX Parse Error: {e}")
        return [], 0


def classify_document_quality(page_count: int, word_count: int, char_count: int, score: int) -> tuple:
    """
    Classifies the document based on extraction metrics.
    Returns (document_quality, richness_reason)
    """
    # 1. Scanned / Extraction Failed
    # If we have many pages but almost no text
    if page_count >= 3 and word_count < 100:
        return "extraction_failed_or_scanned_pdf", f"Document has {page_count} pages but only {word_count} words extracted. Likely a scanned PDF or image-based document."

    # 2. Insufficient Detail (Tiny documents)
    # Very few words regardless of page count (e.g. 1-2 pages with just a header)
    if word_count < 500 and char_count < 3000 and score < 4:
         return "insufficient_rfp_detail", f"Document is too brief ({word_count} words) and lacks standard RFP structure."

    # 3. Valid RFP
    if word_count >= 500 or score >= 4:
        return "valid_rfp", f"Document is sufficient for analysis ({word_count} words, score {score})."
        
    # Fallback
    return "insufficient_rfp_detail", f"Document does not meet the minimum richness threshold (score {score})."

def calculate_document_richness(text: str, sections: list, actual_page_count: int = 0) -> dict:
    """Calculates richness metrics to detect tiny or low-detail RFPs."""
    word_count = len(text.split())
    # Use actual page count if provided, else fallback to section-based estimation
    page_count = actual_page_count or len(set(s.get("page_number", 1) for s in sections))

    
    # Detect requirements and other key segments
    has_requirements = any("requirement" in s["section_name"].lower() for s in sections) or "requirement" in text.lower()
    has_deadline = any(kw in text.lower() for kw in SECTION_KEYWORDS["deadline"])
    has_scope = any(kw in text.lower() for kw in SECTION_KEYWORDS["scope"])
    has_eval_criteria = any("criteria" in s["section_name"].lower() for s in sections) or "criteria" in text.lower()

    # Score calculation (0-10)
    score = 0
    if word_count > 1000: score += 3
    elif word_count > 500: score += 2
    elif word_count > 100: score += 1
    
    if len(sections) > 5: score += 2
    if has_requirements: score += 2
    if has_deadline: score += 1
    if has_scope: score += 1
    if has_eval_criteria: score += 1

    # Scoring
    quality_class, reason = classify_document_quality(page_count, word_count, len(text), score)
    
    return {
        "word_count": word_count,
        "char_count": len(text),
        "page_count": page_count,
        "section_count": len(sections),
        "has_requirements": has_requirements,
        "has_deadline": has_deadline,
        "has_scope": has_scope,
        "has_eval_criteria": has_eval_criteria,
        "richness_score": score,
        "document_quality": quality_class,
        "richness_reason": reason
    }

def run_parsing(db: Session, rfp_id: int):
    """Main function - parse RFP and save sections to DB"""
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        return {"error": "RFP not found"}

    # Update status to parsing
    rfp.current_status = "parsing"
    db.commit()

    try:
        # Parse based on file type
        full_text = ""
        actual_pages = 0
        if rfp.file_type == "pdf":
            sections, actual_pages = parse_pdf(rfp.file_path)
            # Combine text for richness check
            full_text = " ".join(s["section_text"] for s in sections)
        elif rfp.file_type == "docx":
            sections, actual_pages = parse_docx(rfp.file_path)
            full_text = " ".join(s["section_text"] for s in sections)
        else:
            actual_pages = 1
            with open(rfp.file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                full_text = content
                sections = [{
                    "section_name": "Full Document",
                    "section_text": content[:50000],
                    "page_number": 1,
                    "confidence": 100.0,
                    "tags": []
                }]

        # Store full text for persistence and intelligence
        rfp.full_text = full_text

        # Delete old sections if re-parsing
        from app.models.sections import RFPSection
        db.query(RFPSection).filter(RFPSection.rfp_id == rfp_id).delete()

        # Save all sections to DB
        for s in sections:
            section = RFPSection(
                rfp_id=rfp_id,
                section_name=s["section_name"],
                section_text=s["section_text"],
                page_number=s["page_number"],
                confidence=s["confidence"]
            )
            db.add(section)

        # Calculate Richness
        richness = calculate_document_richness(full_text, sections, actual_pages)

        
        # Save to summary_json as a starting point if empty
        if not rfp.summary_json:
            rfp.summary_json = json.dumps(richness)
        else:
            try:
                sj = json.loads(rfp.summary_json)
                sj.update(richness)
                rfp.summary_json = json.dumps(sj)
            except:
                rfp.summary_json = json.dumps(richness)

        # Update status based on quality
        rfp.document_quality = richness["document_quality"]
        rfp.word_count = richness["word_count"]
        rfp.page_count = richness["page_count"]
        rfp.extracted_text_length = richness["char_count"]
        rfp.richness_reason = richness["richness_reason"]
        
        if rfp.document_quality == "valid_rfp":
            rfp.current_status = "parsed"
        elif rfp.document_quality == "extraction_failed_or_scanned_pdf":
            rfp.current_status = "extraction_needs_review"
        else:
            rfp.current_status = "needs_more_detail"

        # Detailed Logging
        print(f"\n[richness_check]")
        print(f"rfp_id={rfp_id}")
        print(f"file_name={rfp.file_name}")
        print(f"file_type={rfp.file_type}")
        print(f"page_count={rfp.page_count}")
        print(f"char_count={rfp.extracted_text_length}")
        print(f"word_count={rfp.word_count}")
        print(f"has_deadline={richness['has_deadline']}")
        print(f"has_scope={richness['has_scope']}")
        print(f"has_requirements={richness['has_requirements']}")
        print(f"document_quality={rfp.document_quality}")
        print(f"current_status={rfp.current_status}")
        print(f"gemini_skipped={rfp.current_status != 'parsed'}")
        print(f"reason={rfp.richness_reason}\n")

        log = AuditLog(rfp_id=rfp_id, action="rfp_parsed",
                       new_value=f"Quality: {rfp.document_quality}. Reason: {rfp.richness_reason}")
        db.add(log)
        db.commit()

        return {
            "status": "parsed",
            "sections_extracted": len(sections),
            "rfp_id": rfp_id,
            "quality": richness["document_quality"]
        }

    except Exception as e:
        import traceback
        print(f"[parsing_service] Error: {e}")
        traceback.print_exc()
        rfp.current_status = "uploaded"
        db.commit()
        return {"error": str(e)}
