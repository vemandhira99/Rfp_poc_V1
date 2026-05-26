from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.api.routes.auth import get_current_user
from app.services import rfp_service
from app.schemas.rfp import (
    RFPOut, DecisionRequest, AssignRequest, CommentRequest,
    DashboardSummary, ChatRequest, DraftRequest, RequirementOut, NotificationOut,
    SubmitReviewRequest, FinalDecisionRequest
)
from datetime import datetime
from typing import List
from app.models.rfp import RFPDocument, RFPDraft, AuditLog

router = APIRouter(prefix="/rfps", tags=["RFPs"])

@router.get("/dashboard-summary")
def dashboard_summary(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    return rfp_service.get_dashboard_summary(db)

@router.get("/", response_model=List[RFPOut])
def list_rfps(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    rfps = rfp_service.get_all_rfps(db)
    return rfps



@router.get("/assigned-to-me", response_model=List[RFPOut])
def assigned_to_me(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    return rfp_service.get_assigned_rfps(db, current_user.id)

@router.get("/notifications", response_model=List[NotificationOut])
def get_notifications(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Returns latest 10 notifications for the current user (read and unread)."""
    from app.services import notification_service
    notifs = notification_service.get_user_notifications(db, current_user.id, limit=10)
    # Debug log to confirm what the API is returning for this user
    print(f"[notification] GET /notifications user_id={current_user.id} role={current_user.role} count={len(notifs)} unread={sum(1 for n in notifs if not n.is_read)}")
    return notifs

@router.get("/notifications/unread-count")
def get_unread_count(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Returns the count of unread notifications for badge display."""
    from app.models.notification import Notification
    count = db.query(Notification).filter(
        Notification.user_id == current_user.id,
        Notification.is_read == False
    ).count()
    return {"unread_count": count, "user_id": current_user.id}

@router.get("/notifications/debug")
def debug_notifications(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Debug endpoint: returns full notification details for the current user."""
    from app.models.notification import Notification
    notifs = db.query(Notification).filter(
        Notification.user_id == current_user.id
    ).order_by(Notification.created_at.desc()).limit(20).all()
    return [
        {
            "id": n.id, "user_id": n.user_id, "rfp_id": n.rfp_id,
            "type": n.type, "message": n.message,
            "is_read": n.is_read, "created_at": str(n.created_at)
        }
        for n in notifs
    ]

@router.post("/notifications/{notification_id}/read")
def mark_notification_read(notification_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    from app.services import notification_service
    notification_service.mark_as_read(db, notification_id)
    return {"message": "Notification marked as read"}

@router.post("/notifications/clear-all")
def clear_all_notifications(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    from app.services import notification_service
    notification_service.clear_all(db, current_user.id)
    return {"message": "All notifications marked as read"}



@router.get("/submitted-for-review", response_model=List[RFPOut])
def get_submitted_rfps_list(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Returns RFPs submitted by SA for PM review."""
    if current_user.role not in ["CEO", "PM", "Admin", "Leadership"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    rfps = db.query(RFPDocument).filter(
        RFPDocument.current_status.in_([
            "submitted_for_review", "revision_requested", "final_approved"
        ])
    ).order_by(RFPDocument.submitted_at.desc()).all()
    return rfps

@router.get("/{rfp_id}", response_model=RFPOut)
def get_rfp(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    rfp = rfp_service.get_rfp_by_id(db, rfp_id)
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    return rfp

@router.post("/{rfp_id}/decision")
def make_decision(rfp_id: int, request: DecisionRequest,
                  db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role not in ["CEO", "Admin", "PM", "Leadership"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    rfp = rfp_service.get_rfp_by_id(db, rfp_id)
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
        
    protected_statuses = [
        "approved", "assigned_to_sa", "generating_draft", "in_drafting", 
        "generation_paused", "draft_complete", "ai_draft_ready", 
        "submitted_for_review", "final_approved"
    ]
    if rfp.current_status in protected_statuses:
        raise HTTPException(status_code=400, detail="Initial PM decision has already been completed for this RFP.")

    result = rfp_service.save_decision(db, rfp_id, current_user.id, request.decision, request.reason)
    return {"message": f"Decision '{request.decision}' saved", "rfp_id": rfp_id}

def background_draft_and_compliance(rfp_id: int, length_mode: str = "short"):
    """Wrapper that calls the new batch-based generation service."""
    from app.services.generation_service import run_generation_job
    run_generation_job(rfp_id, length_mode=length_mode)

@router.post("/{rfp_id}/assign-architect")
def assign_architect(rfp_id: int, request: AssignRequest, background_tasks: BackgroundTasks,
                      db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role not in ["CEO", "PM", "Admin", "Leadership"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    rfp = rfp_service.get_rfp_by_id(db, rfp_id)
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    # Check for existing assignment
    from app.models.rfp import RFPAssignment, RFPDocument
    from app.models.user import User
    existing = db.query(RFPAssignment, User.name).join(User, User.id == RFPAssignment.assigned_to).filter(
        RFPAssignment.rfp_id == rfp_id,
        RFPAssignment.assignment_status == "active"
    ).first()
    
    if existing:
        return {"message": f"RFP is already assigned to {existing.name}.", "rfp_id": rfp_id}

    result = rfp_service.assign_architect(db, rfp_id, request.architect_id, current_user.id, request.notes or "")
    
    # Only trigger generation if not already in progress
    if rfp.current_status not in ["generating_draft", "in_drafting", "draft_complete", "ai_draft_ready"]:
        length_mode = request.length_mode or "short"
        background_tasks.add_task(background_draft_and_compliance, rfp_id, length_mode)
    
    return {"message": "Architect assigned and draft generation started", "rfp_id": rfp_id}
 
@router.post("/{rfp_id}/resume-generation")
def resume_generation(rfp_id: int, background_tasks: BackgroundTasks,
                      db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Resumes a paused or stopped draft generation job."""
    if current_user.role not in ["CEO", "PM", "Admin", "Solution_Architect", "Leadership"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    # Re-trigger generation job; use existing mode
    length_mode = rfp.generation_mode or "short"
    background_tasks.add_task(background_draft_and_compliance, rfp_id, length_mode)
    return {"message": "Generation resumed", "rfp_id": rfp_id}

# ── SA Submit for PM Review ────────────────────────────────────────────────────
@router.post("/{rfp_id}/submit-review")
def submit_for_pm_review(rfp_id: int, request: SubmitReviewRequest,
                         db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    rfp.current_status = "submitted_for_review"
    rfp.submitted_at = datetime.utcnow()
    rfp.submitted_by = current_user.id
    if request.notes:
        rfp.status_message = f"SA Notes: {request.notes}"
    db.add(AuditLog(rfp_id=rfp_id, user_id=current_user.id,
                    action="sa_submitted_for_review", new_value=request.notes or ""))
    db.commit()

    # Notify the PM who assigned this RFP
    try:
        from app.models.rfp import RFPAssignment
        from app.services.notification_service import create_notification
        assignment = db.query(RFPAssignment).filter(
            RFPAssignment.rfp_id == rfp_id,
            RFPAssignment.assignment_status == "active"
        ).first()
        if assignment:
            create_notification(
                db, user_id=assignment.assigned_by,
                message=f"SA has submitted draft for '{rfp.title}' for your review.",
                rfp_id=rfp_id, type="info"
            )
    except Exception as e:
        print(f"[submit_review] Notification failed (non-fatal): {e}")

    return {"message": "Draft submitted for PM review", "status": "submitted_for_review"}


# ── PM Final Approvals / Rejections ───────────────────────────────────────────

@router.post("/{rfp_id}/final-decision")
def final_pm_decision(rfp_id: int, request: FinalDecisionRequest,
                      db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """PM makes a final decision: final_approved | revision_requested | rejected."""
    if current_user.role not in ["CEO", "PM", "Admin", "Leadership"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    valid_decisions = ["final_approved", "revision_requested", "rejected"]
    if request.decision not in valid_decisions:
        raise HTTPException(status_code=400, detail=f"Decision must be one of {valid_decisions}")
    
    rfp.current_status = request.decision
    msg = request.revision_notes or request.reason or ""
    if msg:
        rfp.status_message = f"PM Decision: {msg}"
    db.add(AuditLog(rfp_id=rfp_id, user_id=current_user.id,
                    action=f"pm_final_decision_{request.decision}", new_value=msg))
    db.commit()

    # Notify SA
    try:
        from app.models.rfp import RFPAssignment
        from app.services.notification_service import create_notification
        assignment = db.query(RFPAssignment).filter(
            RFPAssignment.rfp_id == rfp_id,
            RFPAssignment.assignment_status == "active"
        ).first()
        if assignment:
            label = {"final_approved": "Final Approved", "revision_requested": "Revision Requested", "rejected": "Rejected"}.get(request.decision, request.decision)
            create_notification(
                db, user_id=assignment.assigned_to,
                message=f"PM decision for '{rfp.title}': {label}. {msg[:80] if msg else ''}",
                rfp_id=rfp_id, type="success" if request.decision == "final_approved" else "error"
            )
    except Exception as e:
        print(f"[final_decision] Notification failed (non-fatal): {e}")

    return {"message": f"Decision '{request.decision}' recorded", "status": request.decision}


# ── MVP Status Repair ──────────────────────────────────────────────────────────

def derive_mvp_status(rfp) -> str:
    import json
    doc_quality = "sufficient"
    if rfp.summary_json:
        try: doc_quality = json.loads(rfp.summary_json).get("document_quality", "sufficient")
        except: pass

    from app.core.database import SessionLocal
    from app.models.rfp import BackgroundJob, RFPGeneratedDocument, RFPAssignment
    db = SessionLocal()
    gen_doc = db.query(RFPGeneratedDocument).filter(RFPGeneratedDocument.rfp_id == rfp.id).first()
    db.close()

    if doc_quality == "insufficient_rfp_detail":
        return "needs_more_detail"
    
    if rfp.current_status in ["final_approved", "rejected_by_pm", "final_rejected", "revision_requested", "submitted_for_pm_review"]:
        return rfp.current_status
    
    if gen_doc:
        return "ai_draft_ready"
        
    if rfp.current_status in ["generation_running", "generation_queued", "generation_paused"]:
        return rfp.current_status
        
    assignment = db.query(RFPAssignment).filter(RFPAssignment.rfp_id == rfp.id, RFPAssignment.assignment_status == "active").first()
    
    if assignment or rfp.current_status == "assigned_to_architect":
        return "assigned_to_architect"

    if rfp.current_status in ["approved", "approved_by_pm"]:
        return "approved_by_pm"

    if rfp.summary_json:
        return "ready_for_pm_review"
        
    if rfp.current_status in ["analysis_running", "analysis_queued"]:
        return rfp.current_status
        
    return "uploaded"

@router.post("/{rfp_id}/repair-status")
def repair_rfp_status(rfp_id: int, db: Session = Depends(get_db)):
    """Automatically repairs stuck status based on actual artifacts."""
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
        
    new_status = derive_mvp_status(rfp)
    rfp.current_status = new_status
    db.commit()
    return {"message": f"Status repaired to {new_status}", "status": new_status}


# ── Quality Check ─────────────────────────────────────────────────────────────
@router.post("/{rfp_id}/quality-check")
def trigger_quality_check(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Run quality check on draft and return results."""
    from app.services.quality_service import run_quality_check
    result = run_quality_check(db, rfp_id)
    return result

@router.get("/{rfp_id}/quality-status")
def get_quality_status(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Return cached quality_status for the RFP."""
    import json
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    if not rfp.quality_status:
        return {"overall": "not_checked", "detail": "Quality check not yet run"}
    try:
        return json.loads(rfp.quality_status)
    except Exception:
        return {"overall": "error", "detail": "Invalid quality status data"}


@router.post("/{rfp_id}/comment")
def add_comment(rfp_id: int, request: CommentRequest,
                db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    result = rfp_service.add_comment(db, rfp_id, current_user.id,
                                     request.comment_text, request.entity_type, request.entity_id)
    return {"message": "Comment added", "comment_id": result.id}

from app.services import draft_service
from app.schemas.rfp import DraftCreate, DraftOut

@router.post("/{rfp_id}/draft", response_model=DraftOut)
def save_draft(rfp_id: int, request: DraftCreate,
              db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role not in ["Solution_Architect", "CEO", "Admin"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    draft = draft_service.save_draft(db, rfp_id, current_user.id, request.draft_content, request.is_final)
    return draft

@router.get("/{rfp_id}/draft", response_model=DraftOut)
def get_draft(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    draft = draft_service.get_latest_draft(db, rfp_id)
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    return draft

@router.post("/{rfp_id}/chat-stream")
def chat_with_rfp_stream(rfp_id: int, request: ChatRequest, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    try:
        from app.services.ai_service import stream_chat_with_document
        return StreamingResponse(
            stream_chat_with_document(
                db=db, 
                rfp_id=rfp_id, 
                user_id=current_user.id,
                message=request.message, 
                knowledge_mode=request.knowledge_mode,
                history=request.history
            ),
            media_type="text/event-stream"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{rfp_id}/export")
def export_rfp(rfp_id: int, db: Session = Depends(get_db)):
    rfp = rfp_service.get_rfp_by_id(db, rfp_id)
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
    
    # Collect all draft sections for the latest version
    from sqlalchemy import desc
    latest_v_record = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id).order_by(desc(RFPDraft.version)).first()
    
    if not latest_v_record:
        content_sections = [("No Draft Available", "Please generate a draft first.")]
    else:
        version = latest_v_record.version
        drafts = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id, RFPDraft.version == version).order_by(RFPDraft.section_order.asc()).all()
        content_sections = [(d.section_name, d.draft_content) for d in drafts if d.section_name]
    
    import os, time, re
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    from fastapi.responses import FileResponse
    from app.core.config import settings
    from app.models.rfp import RFPRequirement

    doc = Document()
    from app.config.company_profile import get_company_profile
    company = get_company_profile()
    
    # Helper for formatted text
    def add_formatted_text(paragraph, text):
        parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if not part: continue
            if part.startswith('***') and part.endswith('***'):
                run = paragraph.add_run(part[3:-3]); run.bold = True; run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2]); run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1]); run.italic = True
            else:
                paragraph.add_run(part)

    # ─ Confidentiality Notice ─────────────────────────────────────────────────────
    conf_p = doc.add_paragraph(company["confidentiality_text"])
    conf_p.runs[0].font.size = Pt(8)
    conf_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ─ Title Page ───────────────────────────────────────────────────────────
    for _ in range(4): doc.add_paragraph()
    title_p = doc.add_heading(rfp.title, 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2): doc.add_paragraph()
    subtitle = doc.add_paragraph("Comprehensive Technical Proposal Response")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for _ in range(2): doc.add_paragraph()
    client_info = doc.add_paragraph(f"Prepared For:\n{rfp.client_name or 'Issuing Authority'}")
    client_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_info.runs[0].font.size = Pt(12)

    for _ in range(2): doc.add_paragraph()
    submitter = doc.add_paragraph(
        f"Submitted By:\n{company['company_legal_name']}\n"
        f"Platform: {company['product_name']}\n"
        f"{company['website']}"
    )
    submitter.alignment = WD_ALIGN_PARAGRAPH.CENTER
    submitter.runs[0].font.size = Pt(11)

    for _ in range(3): doc.add_paragraph()
    meta_p = doc.add_paragraph(
        f"Version: {latest_v_record.version if latest_v_record else 1}\n"
        f"Date: {time.strftime('%B %d, %Y')}\n"
        f"{company['authorized_signatory']}"
    )
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 2. Table of Contents
    doc.add_heading("1. Table of Contents", level=1)
    for i, (name, _) in enumerate(content_sections):
        doc.add_paragraph(f"1.{i+1} {name}")
    doc.add_page_break()

    # 3. Main Content Sections
    from app.services.generation_service import sanitize_generated_text, strip_duplicate_section_heading
    from app.services.post_processing_service import run_post_processing_pipeline
    
    for i, (name, text) in enumerate(content_sections):
        doc.add_heading(f"{i+1}. {name}", level=1)
        
        if not text: continue
        
        # Apply cleanup services
        text = sanitize_generated_text(text)
        # BUG 5 Fix: Strip duplicate heading
        text = strip_duplicate_section_heading(name, text, i + 1)
        
        # Run the full post-processing pipeline
        text = run_post_processing_pipeline(text)
        
        lines = text.split('\n')
        
        in_table = False
        table_data = []
        
        for line in lines:
            stripped = line.strip()
            
            # Table detection
            if stripped.startswith('|') and stripped.count('|') >= 2:
                if not in_table:
                    in_table = True
                    table_data = []
                
                # Skip separator line (e.g., |---|---|)
                if re.match(r'^\|[\s:-|]*\|$', stripped):
                    continue
                
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if cells:
                    table_data.append(cells)
                continue
            elif in_table:
                # End of table processing
                if table_data:
                    num_cols = max(len(row) for row in table_data)
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(table_data):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx < num_cols:
                                cell_p = table.cell(r_idx, c_idx).paragraphs[0]
                                add_formatted_text(cell_p, cell_text)
                    doc.add_paragraph() # Spacer after table
                in_table = False
                table_data = []

            if not stripped: continue
            
            # Headings within sections
            if stripped.startswith('### '): 
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '): 
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '): 
                doc.add_heading(stripped[2:], level=2) # Treat internal # as H2
            
            # Lists
            elif stripped.startswith('* ') or stripped.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, stripped[2:])
            elif re.match(r'^\d+\.', stripped):
                match = re.match(r'^\d+\.\s*(.*)', stripped)
                if match:
                    p = doc.add_paragraph(style='List Number')
                    add_formatted_text(p, match.group(1))
            
            # Normal paragraph
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped)
        
        # Add page break between major sections to keep it professional
        doc.add_page_break()

    # 4. Compliance Matrix Section (Optional for Short mode)
    if rfp.generation_mode != "short":
        doc.add_heading("Annexure: Compliance Matrix", level=1)
        requirements = db.query(RFPRequirement).filter(RFPRequirement.rfp_id == rfp_id).all()
        if requirements:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'RFP Clause / Section'
            hdr_cells[1].text = 'Requirement Summary'
            hdr_cells[2].text = 'Proposed Compliance'
            hdr_cells[3].text = 'Response Section Reference'
            hdr_cells[4].text = 'Evidence / Remarks'
            hdr_cells[5].text = 'Compliance Status'
            for cell in hdr_cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.bold = True
            for req in requirements:
                row_cells = table.add_row().cells
                row_cells[0].text = str(req.category or "General")
                row_cells[1].text = str(req.requirement_text)
                row_cells[2].text = str(req.status or "Pending")
                row_cells[3].text = str(req.response_strategy or "See Proposal Body")
                row_cells[4].text = str(req.notes or "")
                row_cells[5].text = str(req.status or "Pending")
        
    # 5. Save and Export
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    safe_title = re.sub(r'[^\w\s-]', '', rfp.title).replace(' ', '_')
    file_name = f"Proposal_Draft_{safe_title}_{int(time.time())}.docx"
    file_path = os.path.join(settings.UPLOAD_DIR, file_name)
    doc.save(file_path)
    return FileResponse(file_path, filename=file_name)

@router.get("/{rfp_id}/export-compliance")
def export_compliance_matrix(rfp_id: int, db: Session = Depends(get_db)):
    from docx import Document
    from app.models.rfp import RFPDocument, RFPRequirement, GeneratedDocument
    import os, time, re, logging
    
    logger = logging.getLogger("app")
    try:
        rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
        if not rfp:
            raise HTTPException(status_code=404, detail="RFP not found")
            
        doc = Document()
        doc.add_heading(f"Compliance Matrix - {rfp.title}", level=0)
        
        requirements = db.query(RFPRequirement).filter(RFPRequirement.rfp_id == rfp_id).all()
        if not requirements:
            doc.add_paragraph("No specific requirements found in extraction.")
        else:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'RFP Clause / Section'
            hdr_cells[1].text = 'Requirement Summary'
            hdr_cells[2].text = 'Proposed Compliance'
            hdr_cells[3].text = 'Response Section Reference'
            hdr_cells[4].text = 'Evidence / Remarks'
            hdr_cells[5].text = 'Compliance Status'
            for cell in hdr_cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.bold = True
            for req in requirements:
                row_cells = table.add_row().cells
                row_cells[0].text = str(req.category or "General")
                row_cells[1].text = str(req.requirement_text)
                row_cells[2].text = str(req.status or "Pending")
                row_cells[3].text = str(req.response_strategy or "See Proposal Body")
                row_cells[4].text = str(req.notes or "")
                row_cells[5].text = str(req.status or "Pending")
                
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        safe_title = re.sub(r'[^\w\s-]', '', rfp.title).replace(' ', '_')
        file_name = f"Compliance_Matrix_{safe_title}_{int(time.time())}.docx"
        file_path = os.path.join(settings.UPLOAD_DIR, file_name)
        doc.save(file_path)
        
        # Persist record
        gen_doc = GeneratedDocument(
            rfp_id=rfp_id,
            file_name=file_name,
            file_path=file_path,
            document_type="COMPLIANCE_MATRIX",
            status="ready",
            is_available=True
        )
        db.add(gen_doc)
        db.commit()
        
        return FileResponse(file_path, filename=file_name)
    except Exception as e:
        logger.error(f"Error exporting compliance matrix for RFP {rfp_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate compliance matrix: {str(e)}")

import json
from app.services.ai_service import chat_with_document

@router.get("/{rfp_id}/summary")
def get_ai_summary(rfp_id: int, db: Session = Depends(get_db)):
    log = db.query(AuditLog).filter(AuditLog.rfp_id == rfp_id, AuditLog.action == "ai_summary_generated").order_by(AuditLog.created_at.desc()).first()
    if not log or not log.new_value: return {"error": "Summary not found"}
    try: return json.loads(log.new_value)
    except: return {"error": "Invalid summary format"}

@router.post("/{rfp_id}/chat")
def chat_with_rfp(rfp_id: int, request: ChatRequest, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    try:
        reply_dict = chat_with_document(db=db, rfp_id=rfp_id, user_id=current_user.id, message=request.message, knowledge_mode=request.knowledge_mode, history=request.history)
        return reply_dict
    except Exception as e:
        return {"reply": f"AI Advisor is currently unavailable. ({str(e)})", "metadata": {"error": True}}

@router.get("/{rfp_id}/compliance", response_model=List[RequirementOut])
def get_compliance_matrix(rfp_id: int, db: Session = Depends(get_db)):
    from app.models.rfp import RFPRequirement
    return db.query(RFPRequirement).filter(RFPRequirement.rfp_id == rfp_id).all()


@router.post("/{rfp_id}/regenerate")
def regenerate_proposal(rfp_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    from app.services import rfp_service
    rfp_service.update_rfp_status(db, rfp_id, "generating_draft", current_user.id)
    background_tasks.add_task(background_draft_and_compliance, rfp_id, "short")
    return {"message": "Regeneration started", "status": "processing"}

@router.post("/{rfp_id}/cancel-generation")
def cancel_generation(rfp_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    from app.services import rfp_service
    rfp_service.update_rfp_status(db, rfp_id, "on_hold", current_user.id)
    return {"message": "Generation cancellation requested", "status": "cancelling"}

@router.get("/{rfp_id}/generation-progress")
def get_generation_progress(rfp_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    from sqlalchemy import desc
    from app.utils.status_map import get_display_label
    from app.services.generation_service import get_sections_for_mode

    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    gen_mode = rfp.generation_mode if rfp and rfp.generation_mode else "short"
    mode_sections = get_sections_for_mode(gen_mode)
    total_sections = len(mode_sections)

    latest_v_record = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id).order_by(desc(RFPDraft.version)).first()
    
    if not latest_v_record:
        return {"current": 0, "total": total_sections, "status": "starting", "display_status": "Draft Queued", "progress_percentage": 0, "is_complete": False}
    
    latest_version = latest_v_record.version
    count = db.query(RFPDraft).filter(
        RFPDraft.rfp_id == rfp_id,
        RFPDraft.version == latest_version,
        RFPDraft.section_name != None
    ).count()
    
    from app.models.rfp import GeneratedDocument
    # Get latest proposal draft
    generated_doc = db.query(GeneratedDocument).filter(
        GeneratedDocument.rfp_id == rfp_id,
        GeneratedDocument.document_type == "PROPOSAL_DRAFT"
    ).order_by(desc(GeneratedDocument.version)).first()

    # Get latest compliance matrix
    compliance_doc = db.query(GeneratedDocument).filter(
        GeneratedDocument.rfp_id == rfp_id,
        GeneratedDocument.document_type == "COMPLIANCE_MATRIX"
    ).order_by(desc(GeneratedDocument.version)).first()

    current_status = rfp.current_status if rfp else "unknown"
    is_complete = (generated_doc is not None) or (count >= total_sections) or \
                  (current_status in ["draft_complete", "ai_draft_ready", "submitted_for_review", "final_approved"])
    
    if is_complete:
        status = "completed"
        if current_status in ["generating_draft", "in_drafting", "processing"]:
             rfp.current_status = "ai_draft_ready"
             db.commit()
    elif current_status in ["on_hold", "on-hold", "cancelled"]:
        status = "cancelled"
    elif current_status in ["error", "rate_limit_error"]:
        status = "failed"
    else:
        status = "running"
        
    progress_pct = round((count / total_sections) * 100) if total_sections > 0 else 0
    if is_complete: progress_pct = 100

    return {
        "current": count,
        "total": total_sections,
        "progress_percentage": progress_pct,
        "status": status,
        "display_status": get_display_label(current_status, progress_pct),
        "is_complete": is_complete,
        "status_message": rfp.status_message if rfp else "",
        "generation_mode": gen_mode,
        "last_activity": rfp.updated_at.isoformat() if rfp and rfp.updated_at else None,
        "generated_document": {
            "id": generated_doc.id,
            "file_name": generated_doc.file_name,
            "version": generated_doc.version,
            "created_at": generated_doc.created_at.isoformat()
        } if generated_doc else None,
        "compliance_matrix": {
            "id": compliance_doc.id,
            "file_name": compliance_doc.file_name,
            "version": compliance_doc.version,
            "created_at": compliance_doc.created_at.isoformat(),
            "is_available": compliance_doc.is_available
        } if compliance_doc else None
    }

@router.post("/{rfp_id}/export-compliance-matrix")
def create_compliance_matrix(rfp_id: int, db: Session = Depends(get_db)):
    from docx import Document
    from app.models.rfp import RFPDocument, RFPRequirement, GeneratedDocument
    import os, time, re
    
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")
        
    doc = Document()
    doc.add_heading(f"Compliance Matrix - {rfp.title}", level=0)
    
    requirements = db.query(RFPRequirement).filter(RFPRequirement.rfp_id == rfp_id).all()
    if not requirements:
        doc.add_paragraph("No specific requirements found in extraction.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'RFP Clause / Section'
        hdr_cells[1].text = 'Requirement Summary'
        hdr_cells[2].text = 'Proposed Compliance'
        hdr_cells[3].text = 'Response Section Reference'
        hdr_cells[4].text = 'Evidence / Remarks'
        hdr_cells[5].text = 'Compliance Status'
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs: r.bold = True
        for req in requirements:
            row_cells = table.add_row().cells
            row_cells[0].text = str(req.category or "General")
            row_cells[1].text = str(req.requirement_text)
            row_cells[2].text = str(req.status or "Pending")
            row_cells[3].text = str(req.response_strategy or "See Proposal Body")
            row_cells[4].text = str(req.notes or "")
            row_cells[5].text = str(req.status or "Pending")
            
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    safe_title = re.sub(r'[^\w\s-]', '', rfp.title).replace(' ', '_')
    file_name = f"Compliance_Matrix_{safe_title}_{int(time.time())}.docx"
    file_path = os.path.join(settings.UPLOAD_DIR, file_name)
    doc.save(file_path)
    
    # Persist in GeneratedDocument
    gen_doc = GeneratedDocument(
        rfp_id=rfp_id,
        file_name=file_name,
        file_path=file_path,
        document_type="COMPLIANCE_MATRIX",
        status="ready",
        is_available=True
    )
    db.add(gen_doc)
    db.commit()
    db.refresh(gen_doc)
    
    return {"message": "Compliance matrix generated", "file_name": file_name, "id": gen_doc.id}

@router.get("/{rfp_id}/download-generated/{file_name}")
def download_generated_doc(rfp_id: int, file_name: str, db: Session = Depends(get_db)):
    from app.models.rfp import GeneratedDocument
    from fastapi.responses import FileResponse
    import os, logging
    
    logger = logging.getLogger("app")
    logger.info(f"Download request: RFP {rfp_id}, File {file_name}")
    
    if not file_name or file_name == "undefined" or file_name == "null":
        raise HTTPException(status_code=400, detail="Compliance matrix is not available for this RFP.")
    
    try:
        doc = db.query(GeneratedDocument).filter(
            GeneratedDocument.rfp_id == rfp_id,
            GeneratedDocument.file_name == file_name
        ).first()
        
        if not doc:
            logger.warning(f"Document record not found in DB: {file_name}")
            raise HTTPException(status_code=400, detail="Compliance matrix is not available for this RFP.")
            
        if not os.path.exists(doc.file_path):
            logger.error(f"File missing on disk at {doc.file_path}")
            raise HTTPException(status_code=400, detail="Compliance matrix is not available for this RFP.")
            
        logger.info(f"Serving file: {doc.file_path}")
        return FileResponse(doc.file_path, filename=doc.file_name)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading generated doc: {str(e)}")
        raise HTTPException(status_code=400, detail="Compliance matrix is not available for this RFP.")


# ════════════════════════════════════════════════════════════════════════════════
# PART 5: Job Progress & Stuck Job Recovery Endpoints
# ════════════════════════════════════════════════════════════════════════════════

@router.get("/{rfp_id}/analysis-progress")
def get_analysis_progress(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Get the current progress of an analysis job for an RFP."""
    from app.services.job_service import get_job_progress
    return get_job_progress(db, rfp_id, "analysis")


@router.get("/{rfp_id}/generation-progress")
def get_generation_progress(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Get the current progress of a generation job for an RFP."""
    from app.services.job_service import get_job_progress
    return get_job_progress(db, rfp_id, "generation")


@router.post("/{rfp_id}/retry-analysis")
def retry_analysis(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Retry a failed analysis job. Resets status and re-queues."""
    from app.models.rfp import RFPDocument, AuditLog
    from app.services.job_service import create_job

    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    job = create_job(db, rfp_id, "analysis", total_steps=5)
    rfp.current_status = "queued_for_processing"
    rfp.status_message = "Retry requested by user."
    db.add(AuditLog(rfp_id=rfp_id, user_id=current_user.id, action="analysis_retry_requested"))
    db.commit()

    # Trigger background analysis
    from app.services.ai_service import analyze_rfp_with_gemini
    import asyncio
    import threading
    t = threading.Thread(target=analyze_rfp_with_gemini, args=(rfp_id, db.__class__()))
    t.daemon = True
    t.start()
    return {"job_id": job.id, "status": "queued", "message": "Analysis re-queued successfully."}


@router.post("/{rfp_id}/resume-generation")
def resume_generation(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Resume a paused generation job from where it stopped."""
    from app.models.rfp import RFPDocument, AuditLog, RFPDraft
    from app.services.job_service import get_job_progress

    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    if rfp.current_status not in ["paused", "failed_retryable", "generation_paused"]:
        return {"message": f"RFP is in '{rfp.current_status}' — no resume needed.", "status": rfp.current_status}

    rfp.current_status = "generating_draft"
    db.add(AuditLog(rfp_id=rfp_id, user_id=current_user.id, action="generation_resumed"))
    db.commit()
    return {"message": "Generation resumed.", "status": "generating_draft"}


@router.post("/{rfp_id}/repair-status")
def repair_rfp_status(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    Intelligent status repair. Infers correct status from DB evidence.
    Use when RFP is stuck in 'AI Analyzing', 'rate_limit_error', or other stale states.
    """
    from app.services.job_service import repair_rfp_status as _repair
    result = _repair(rfp_id, db, commit=True)
    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])
    return result


@router.post("/repair-all-stuck")
def repair_all_stuck_rfps(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Admin endpoint: scan and repair all stuck jobs and orphaned RFP statuses."""
    if current_user.role not in ["Admin", "CEO", "PM"]:
        raise HTTPException(status_code=403, detail="Permission denied")
    from app.services.job_service import detect_and_repair_stuck_jobs
    repaired = detect_and_repair_stuck_jobs(db)
    return {"repaired_count": len(repaired), "items": repaired}


# ════════════════════════════════════════════════════════════════════════════════
# PART 3: Human Approval Gate Endpoints
# ════════════════════════════════════════════════════════════════════════════════

@router.get("/{rfp_id}/generation-confirmation-data")
def get_generation_confirmation(rfp_id: int, mode: str = "standard", db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    Returns data needed to populate the generation confirmation modal.
    PM/CEO must review this before starting draft generation.
    """
    from app.services.job_service import get_generation_confirmation_data
    return get_generation_confirmation_data(rfp_id, current_user, mode, db)


@router.post("/{rfp_id}/confirm-generation")
def confirm_and_start_generation(rfp_id: int, request: dict, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    PART 3: Human Approval Gate for Draft Generation.
    
    PM/SA must explicitly confirm before generation starts.
    Request body: { "mode": "standard", "architect_id": 123, "confirmed": true }
    """
    if not request.get("confirmed"):
        raise HTTPException(status_code=400, detail="Generation requires explicit confirmation. Set 'confirmed': true.")

    mode = request.get("mode", "standard")
    architect_id = request.get("architect_id")

    from app.models.rfp import RFPDocument, AuditLog, RFPDraft
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    # Guard: prevent if already generating
    if rfp.current_status in ["generating_draft", "in_drafting"]:
        raise HTTPException(status_code=409, detail=f"Generation already in progress. Status: {rfp.current_status}")

    # Guard: require approved/assigned status
    allowed_statuses = ["approved", "assigned_to_sa", "ai_draft_ready", "under_sa_review", "pending-review"]
    if rfp.current_status not in allowed_statuses:
        raise HTTPException(status_code=400, detail=f"Cannot start generation from status '{rfp.current_status}'.")

    rfp.current_status = "generating_draft"
    rfp.proposal_length_mode = mode
    db.add(AuditLog(
        rfp_id=rfp_id,
        user_id=current_user.id,
        action="generation_confirmed",
        new_value=f"mode={mode}, architect_id={architect_id}",
    ))
    db.commit()

    # Kick off background generation
    import threading
    def _run_generation():
        from app.core.database import SessionLocal
        gen_db = SessionLocal()
        try:
            from app.services.generation_service import generate_full_proposal
            generate_full_proposal(gen_db, rfp_id, mode)
        except Exception as e:
            print(f"[confirm-generation] Generation error for RFP {rfp_id}: {e}")
        finally:
            gen_db.close()

    t = threading.Thread(target=_run_generation, daemon=True)
    t.start()

    return {"message": "Draft generation started.", "mode": mode, "rfp_id": rfp_id, "status": "generating_draft"}


@router.post("/{rfp_id}/confirm-regeneration")
def confirm_regeneration(rfp_id: int, request: dict, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    PART 3: Regeneration requires explicit overwrite confirmation.
    Request body: { "confirmed_overwrite": true, "mode": "standard" }
    """
    if not request.get("confirmed_overwrite"):
        raise HTTPException(
            status_code=400,
            detail="Regeneration requires explicit overwrite confirmation. Existing sections will be deleted. Set 'confirmed_overwrite': true."
        )

    from app.models.rfp import RFPDocument, RFPDraft, AuditLog
    rfp = db.query(RFPDocument).filter(RFPDocument.id == rfp_id).first()
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    mode = request.get("mode", rfp.proposal_length_mode or "standard")

    # Delete existing draft sections
    deleted = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id).delete()
    rfp.current_status = "generating_draft"
    rfp.proposal_length_mode = mode
    db.add(AuditLog(
        rfp_id=rfp_id,
        user_id=current_user.id,
        action="regeneration_confirmed",
        new_value=f"mode={mode}, sections_deleted={deleted}",
    ))
    db.commit()

    import threading
    def _run_generation():
        from app.core.database import SessionLocal
        gen_db = SessionLocal()
        try:
            from app.services.generation_service import generate_full_proposal
            generate_full_proposal(gen_db, rfp_id, mode)
        except Exception as e:
            print(f"[confirm-regeneration] Error for RFP {rfp_id}: {e}")
        finally:
            gen_db.close()

    threading.Thread(target=_run_generation, daemon=True).start()
    return {"message": f"Regeneration started. {deleted} existing sections removed.", "rfp_id": rfp_id, "mode": mode}


# ════════════════════════════════════════════════════════════════════════════════
# PART 1: Context & Gemini URI Endpoints
# ════════════════════════════════════════════════════════════════════════════════

@router.get("/{rfp_id}/context")
def get_rfp_context(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    Return full persisted context for an RFP from PostgreSQL only.
    No AI call made. Safe to call after 2-3 week PM absence.
    """
    from app.services.context_service import load_rfp_context
    ctx = load_rfp_context(rfp_id, db)
    if not ctx:
        raise HTTPException(status_code=404, detail="RFP not found")
    return ctx


@router.post("/{rfp_id}/ensure-gemini-file")
def ensure_gemini_file(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    Re-upload original file to Gemini if URI has expired.
    Only call this when deep document reasoning is required.
    """
    from app.services.context_service import ensure_gemini_file_available
    uri = ensure_gemini_file_available(rfp_id, db)
    return {
        "rfp_id": rfp_id,
        "gemini_uri": uri,
        "available": uri is not None,
        "message": "Gemini file ready." if uri else "Original file not found or upload failed."
    }


@router.get("/{rfp_id}/evaluation-summary")
def get_evaluation_summary(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Return evaluator status summary across all draft sections."""
    from app.models.rfp import RFPDraft
    from app.services.context_service import _build_evaluation_summary
    drafts = db.query(RFPDraft).filter(RFPDraft.rfp_id == rfp_id).all()
    return {"rfp_id": rfp_id, **_build_evaluation_summary(drafts)}


# BUG 4 FIX: Section Order Repair Endpoint
@router.post("/{rfp_id}/repair-section-order")
def repair_section_order_endpoint(rfp_id: int, db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """
    BUG 4 FIX: Repair section_order for any existing draft records with NULL or wrong values.
    Maps known section titles to their canonical order number from the proposal table of contents.
    Safe to call anytime — only updates records that differ from canonical order.
    """
    from app.services.generation_service import repair_section_order
    repaired = repair_section_order(db, rfp_id)
    return {
        "rfp_id": rfp_id,
        "repaired_count": repaired,
        "message": f"{repaired} sections had their order fixed." if repaired else "All sections already have correct order."
    }


@router.post("/repair-all-section-orders")
def repair_all_section_orders(db: Session = Depends(get_db), current_user=Depends(get_current_user)):
    """Admin: repair section_order for ALL RFPs with generated drafts."""
    if current_user.role not in ["Admin", "CEO", "PM"]:
        raise HTTPException(status_code=403, detail="Permission denied")
    from app.models.rfp import RFPDocument, RFPDraft
    from app.services.generation_service import repair_section_order
    rfp_ids = [r[0] for r in db.query(RFPDraft.rfp_id).distinct().all()]
    total_repaired = 0
    for rfp_id in rfp_ids:
        total_repaired += repair_section_order(db, rfp_id)
    return {"rfps_processed": len(rfp_ids), "total_sections_repaired": total_repaired}


